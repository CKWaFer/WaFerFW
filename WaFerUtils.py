import os
import pandas as pd
import numpy as np
import torch
from natsort import ns, natsorted
import json
import math
# import SimpleITK as sitk
import shutil
from scipy import stats
import warnings
warnings.filterwarnings('ignore')

def predo_root(root: str):
    root = root.replace('\\', '/')
    return root if root[-1] == '/' else root + '/'


def make_folder(root: str):
    root = predo_root(root)
    if os.path.exists(root):
        if root[-1] == '/':
            root = root[:-1]
        root += '_0/'
        make_folder(root)
    else:
        os.makedirs(root)
    pass


def dpabi_workspace(cur_root: str,
                    save_root: str,
                    cur_layers_num=0,  # 当前递归层数
                    cur_types_num=0,  # 当前文件序号
                    rate=0.9,  # 训练集占比
                    train_or_test='train'  # 训练集/测试集 标识符
                    ):
    # 规范路径名格式
    cur_root = predo_root(cur_root)
    save_root = predo_root(save_root)

    if not os.path.exists(save_root):
        os.makedirs(save_root)

    cur_dir = os.listdir(cur_root)  # 获取当前文件夹下的文件名
    cur_dir_len = len(cur_dir)  # 文件数

    for i in range(cur_dir_len):
        if cur_layers_num == 0:
            train_or_test = 'train' if i <= (cur_dir_len * rate) else 'lyj'
        fn = cur_dir[i]  # 当前文件/文件夹名

        # 拼接路径名
        new_path = cur_root + fn  # 可能是文件，末尾不能拼接'/'
        if os.path.isdir(new_path):  # 是否为文件夹
            new_path = predo_root(new_path)  # 新文件夹路径
            # 进入递归
            dpabi_workspace(cur_root=new_path,
                            save_root=save_root,
                            cur_layers_num=cur_layers_num + 1,
                            cur_types_num=i,
                            train_or_test=train_or_test)
        else:
            # 发现了文件
            file_path = cur_root + fn
            print("\r> 正在获取文件路径：{}".format(fn), end='')
            cur_save_path = save_root + train_or_test + '-' + str(cur_layers_num) + '-' + str(cur_types_num) + '.txt'

            with open(cur_save_path, mode='a') as f:  # 打开文件
                f.write("{}\n".format(file_path))
                f.close()


# dpabi_workspace(cur_root='/home/wafer/WorkSpace/DataSet/tdt',
#                 save_root='./lyj/')


# 本函数通过CSV文件中的信息，可以将被试划分成训练集和测试集，并把被试相关信息和文件路径封装到json文件中
def create_train_val_by_data_and_csv_json(data_root: str,
                                          csv_path: str,
                                          save_root: str,
                                          csv_id_keyword: str,
                                          csv_keywords: list,
                                          train_rate=0.7,
                                          shuffle=True,
                                          csv_file_type='.csv'
                                         ):
    data_root = predo_root(data_root)
    save_root = predo_root(save_root)

    if os.path.exists(save_root):
        return
        pass

    csv_data = None
    dir_files = None
    if csv_path is not None:
        csv_data = pd.read_csv(csv_path)
    if os.path.isdir(data_root):
        dir_files = os.listdir(data_root)
        if len(dir_files) <= 0:
            return None

    match_files = list()
    id_list = [str(i).split('.')[0] for i in csv_data[csv_id_keyword]]
    id_list = list(filter(lambda x: x != 'nan', id_list))
    cur_id_dict = dict()

    for f in dir_files:
        # 这里是以现有文件去校正从csv文件获取的id列表，下面内容为补丁
        id = list(filter(lambda x: f.find(x) >= 0, id_list))
        if len(id) > 0:
            cur_id_dict[id[0]] = None
        if '.mat' in f or '.nii' in f or '.1D' in f:
            flag = False
            for b in id_list:
                if f.find(b) >= 0:
                    flag = True
                    break
            if flag is True:
                match_files.append(f)

    id_list = cur_id_dict.keys()

    dir_files = match_files
    dir_files = natsorted(dir_files, alg=ns.PATH)

    # 被试列表，包含被试的id，标签和所有文件路径
    bs_list = list()
    for id in id_list:
        bs_dict = dict()  # 被试字典
        bs_dict['id'] = str(id)  # 存入被试id

        # 通过关键字，在csv中找到指定属性
        for ck in csv_keywords:
            if ck == 'DX_GROUP':  # 组别调整为 0~1  (原始为1~2)
                v = str(int(csv_data[csv_data[csv_id_keyword] == int(id)][ck]) - 1)
            else:
                v = csv_data[csv_data[csv_id_keyword] == int(id)][ck]
            if type(v) is pd.core.series.Series:
                v = str(v.values[0])
            bs_dict[ck] = v

        fp_idx = 0  # 文件下标
        for df in dir_files:
            if df.find(id) >= 0:  # 文件名里包含id
                bs_dict['file{}'.format(fp_idx)] = '{}{}'.format(data_root, df)  # 将文件存入字典
                fp_idx += 1
        bs_list.append(bs_dict)

    if shuffle is True:
        shuffle_list = list()
        order = np.random.permutation(np.array(range(0, len(bs_list), 1)))
        for i in order:
            b = bs_list[i]
            shuffle_list.append(b)
        bs_list = shuffle_list

    train_num = math.ceil(len(bs_list) * train_rate)
    val_num = int(len(bs_list) - train_num)

    bs_train_list = [bs_list[i] for i in range(train_num)]
    bs_val_list = [bs_list[train_num + i] for i in range(val_num)]

    if not os.path.exists(save_root):
        os.makedirs(save_root)

    train_save_root = save_root + 'train/'
    if not os.path.exists(train_save_root):
        os.mkdir(train_save_root)


    # 写入训练集
    for bs in bs_train_list:
        str_id = bs['id']

        path = train_save_root + str_id + '.json'
        with open(path, "w", encoding='utf-8') as f:
            json.dump(bs, f, indent=2, sort_keys=True, ensure_ascii=False)  # 写为多行

    val_save_root = save_root + 'val/'
    if not os.path.exists(val_save_root):
        os.mkdir(val_save_root)

    # 写入测试集
    for bs in bs_val_list:
        str_id = bs['id']

        path = val_save_root + str_id + '.json'
        with open(path, "w", encoding='utf-8') as f:
            json.dump(bs, f, indent=2, sort_keys=True, ensure_ascii=False)  # 写为多行


# 本函数通过CSV文件中的信息，可以将被试划分成训练集和测试集，并把被试相关信息和文件路径封装到json文件中
def getDataByCsvJson(data_root: str,
                     csv_path: str,
                     save_root: str,
                     csv_id_keyword: str,
                     csv_keywords: list,
                     train_rate=0.7,
                     shuffle=True,
                     csv_file_type='.csv'
                     ):
    data_root = predo_root(data_root)
    save_root = predo_root(save_root)

    if os.path.exists(save_root):
        return
        pass

    csv_data = None
    dir_files = None
    if csv_path is not None:
        csv_data = pd.read_csv(csv_path)
    if os.path.isdir(data_root):
        dir_files = os.listdir(data_root)
        if len(dir_files) <= 0:
            return None

    match_files = list()
    id_list = [str(i).split('.')[0] for i in csv_data[csv_id_keyword]]
    id_list = list(filter(lambda x: x != 'nan', id_list))
    cur_id_dict = dict()

    for f in dir_files:
        # 这里是以现有文件去校正从csv文件获取的id列表，下面内容为补丁
        id = list(filter(lambda x: f.find(x) >= 0, id_list))
        if len(id) > 0:
            cur_id_dict[id[0]] = None
        if '.mat' in f or '.nii' in f or '.1D' in f:
            flag = False
            for b in id_list:
                if f.find(b) >= 0:
                    flag = True
                    break
            if flag is True:
                match_files.append(f)

    id_list = cur_id_dict.keys()

    dir_files = match_files
    dir_files = natsorted(dir_files, alg=ns.PATH)

    # 被试列表，包含被试的id，标签和所有文件路径
    bs_list = list()
    for id in id_list:
        bs_dict = dict()  # 被试字典
        bs_dict['id'] = str(id)  # 存入被试id

        # 通过关键字，在csv中找到指定属性
        for ck in csv_keywords:
            if ck == 'DX_GROUP':  # 组别调整为 0~1  (原始为1~2)
                v = str(int(csv_data[csv_data[csv_id_keyword] == int(id)][ck]) - 1)
            else:
                v = csv_data[csv_data[csv_id_keyword] == int(id)][ck]
            if type(v) is pd.core.series.Series:
                v = str(v.values[0])
            bs_dict[ck] = v

        fp_idx = 0  # 文件下标
        for df in dir_files:
            if df.find(id) >= 0:  # 文件名里包含id
                bs_dict['file{}'.format(fp_idx)] = '{}{}'.format(data_root, df)  # 将文件存入字典
                fp_idx += 1
        bs_list.append(bs_dict)

    if shuffle is True:
        shuffle_list = list()
        order = np.random.permutation(np.array(range(0, len(bs_list), 1)))
        for i in order:
            b = bs_list[i]
            shuffle_list.append(b)
        bs_list = shuffle_list

    train_num = math.ceil(len(bs_list) * train_rate)
    val_num = int(len(bs_list) - train_num)

    bs_train_list = [bs_list[i] for i in range(train_num)]
    bs_val_list = [bs_list[train_num + i] for i in range(val_num)]

    if not os.path.exists(save_root):
        os.makedirs(save_root)

    train_save_root = save_root + 'train/'
    if not os.path.exists(train_save_root):
        os.mkdir(train_save_root)


    # 写入训练集
    for bs in bs_train_list:
        str_id = bs['id']

        path = train_save_root + str_id + '.json'
        with open(path, "w", encoding='utf-8') as f:
            json.dump(bs, f, indent=2, sort_keys=True, ensure_ascii=False)  # 写为多行

    val_save_root = save_root + 'val/'
    if not os.path.exists(val_save_root):
        os.mkdir(val_save_root)

    # 写入测试集
    for bs in bs_val_list:
        str_id = bs['id']

        path = val_save_root + str_id + '.json'
        with open(path, "w", encoding='utf-8') as f:
            json.dump(bs, f, indent=2, sort_keys=True, ensure_ascii=False)  # 写为多行




# 将子文件夹文件拿到指定位置
def get_data_in_one_folder(root: str, save_root: str, n=0):
    root = std_root(root)
    save_root = std_root(save_root)

    if not os.path.exists(save_root):
        os.makedirs(save_root)

    if (n == 0) and (len(os.listdir(save_root)) > 0):
        return
        pass

    dirs = get_dirs(root)

    for dir in dirs:
        cur_root = std_root(root + dir)
        get_data_in_one_folder(root=cur_root, save_root=save_root, n=n)
        n += 1

    fns = get_fns(root)
    for fn in fns:
        cur_path = root + fn
        new_fn = root.split('/')[-2] + "_" + fn
        new_path = save_root + new_fn
        if os.path.isfile(cur_path):
            if not os.path.exists(save_root):
                os.makedirs(save_root)
            print("\r正在复制: {}".format(new_fn), end='')
            shutil.copyfile(cur_path, new_path)


def json_reader(json_path: str,
                key_words: list
                ):
    output = dict()
    with open(json_path, 'r', encoding='utf-8') as f:
        content = f.read()
        json_data = json.loads(content)
        for kw in key_words:
            info = ''
            if kw in json_data.keys():
                info = json_data[kw]
            output[kw] = info
        f.close()

    return output


def get_bs_info_from_json(json_root: str,
                          key_words: list
                          ):
    json_root = predo_root(json_root)
    dir_files = None
    if os.path.exists(json_root):
        dir_files = os.listdir(json_root)
    dir_files = natsorted(dir_files, alg=ns.PATH)  # 排序

    bs_info_list = list()
    for df in dir_files:
        path = json_root + df
        if os.path.exists(path):
            bs_info = json_reader(json_path=path, key_words=key_words)
            bs_info_list.append(bs_info)

    return bs_info_list


# 按照指定 维度和步长 对4D的fMRI切片(步长默认为1)
def fmir_4d_to_3d(data: np.ndarray, axis=None, length=1):
    res_list = list()
    # 如果传入指定维度(axis is not None)，则按照指定维度切割，返回切割结果
    # 如果没有传入指定维度(axis is None)，则将所有维度都做一次切割，返回所有的切割结果
    if axis is not None:
        num = math.ceil(data.shape[axis] / length)  # 计算要切分的份数
        # 返回分割好的数据，并将其维度降低1个维度
        return [np.squeeze(d, axis=axis) for d in np.split(ary=data, indices_or_sections=num, axis=axis)]
    else:
        for axis in range(len(data.shape)):
            num = math.ceil(data.shape[axis] / length)  # 计算要切分的份数
            # 返回分割好的数据，并将其维度降低1个维度
            res_list.append([np.squeeze(d, axis=axis) for d in np.split(ary=data, indices_or_sections=num, axis=axis)])
    return res_list


# 将3D的fMRI按照时间维度切割(我们认为最后一个维度是时间维度)
def sliding_3d_fmir(data: np.ndarray, t: int):
    res_list = list()
    tp_len = data.shape[-1]  # 时间点长度
    start_idx = tp_len % t   # 由于保证每段一样长，这里将多出的时间点从开始处去除

    # 按照时间维度，间隔t，切分4DfMRI
    for i in range(start_idx, tp_len, t):
        res_list.append(data[:, :, i: i + t])

    return res_list


# 将4D的fMRI按照时间维度切割(我们认为最后一个维度是时间维度)
def sliding_4d_fmir(data: np.ndarray, t: int):
    res_list = list()
    tp_len = data.shape[-1]  # 时间点长度
    start_idx = tp_len % t   # 由于保证每段一样长，这里将多出的时间点从开始处去除

    # 按照时间维度，间隔t，切分4DfMRI
    for i in range(start_idx, tp_len, t):
        res_list.append(data[:, :, :, i: i + t])

    return res_list


# 路径格式标准化
def std_root(root: str):
    root = root.replace('\\', '/')
    root = root if root[-1] == '/' else root + '/'

    return root


# 寻找当前路径下的文件夹
def get_dirs(root: str):
    res = list()
    dir_files = os.listdir(root)
    dir_files = natsorted(dir_files, alg=ns.PATH)  # 此处需要 from natsort import ns, natsorted
    for fn in dir_files:
        if (fn[0] == '.') or os.path.isfile(root + fn):
            continue
        res.append(fn)

    return res


# 寻找当前路径下的文件
def get_fns(root: str):
    res = list()
    dir_files = os.listdir(root)
    dir_files = natsorted(dir_files, alg=ns.PATH)  # 此处需要 from natsort import ns, natsorted
    for fn in dir_files:
        if (fn[0] == '.') or os.path.isdir(root + fn):
            continue
        res.append(fn)

    return res


def read_bs_csv(path: str):
    csv_data = pd.read_csv(path)
    csv_keys = csv_data.keys()
    bsid_list = csv_data['SUB_ID'].tolist()
    new_bsid_list = list(set(bsid_list))
    new_bsid_list.sort(key=bsid_list.index)
    bsid_list = new_bsid_list

    return csv_data, csv_keys, bsid_list


# # 改变 MRI尺寸
# # 代码来源：https://blog.csdn.net/jancis/article/details/106265602
# def resize_image_itk(itkimage, newSize, resamplemethod=sitk.sitkNearestNeighbor):
#     resampler = sitk.ResampleImageFilter()
#     originSize = itkimage.GetSize()  # 原来的体素块尺寸
#     originSpacing = itkimage.GetSpacing()
#     newSize = np.array(newSize, float)
#     factor = originSize / newSize
#     newSpacing = originSpacing * factor
#     newSize = newSize.astype(np.int)  # spacing肯定不能是整数
#     resampler.SetReferenceImage(itkimage)  # 需要重新采样的目标图像
#     resampler.SetSize(newSize.tolist())
#     resampler.SetOutputSpacing(newSpacing.tolist())
#     resampler.SetTransform(sitk.Transform(3, sitk.sitkIdentity))
#     resampler.SetInterpolator(resamplemethod)
#     itkimgResampled = resampler.Execute(itkimage)  # 得到重新采样后的图像
#     return itkimgResampled


# 读取 脚本提取的预处理的数据
def read_predo_data(json_root: str,
                    new_size=None,  # 如果要改变尺寸，传入(x, x, x)  [仅三维，时间维度不用写，也不能改]
                    time_points_num=None,  # 要读取的时间点，默认为全读取
                    is_train=None
                    ):
    json_root = predo_root(json_root)

    bss_info = get_bs_info_from_json(json_root=json_root, key_words=['id', 'DX_GROUP', 'file0'])

    bs_list = list()
    label_list = list()
    c0 = 0
    for bs in bss_info:
        id = bs['id']
        y = bs['DX_GROUP']
        fp = bs['file0']
        fn = fp.split('/')[-1]

        c0 += 1

        mode = ""
        if is_train is not None:
            mode = "<训练集>" if is_train is True else "<验证集>"
        print("\r({}/{}) - 正在读取{}: {}".format(c0, len(bss_info), mode, fn), end='')

        img = sitk.ReadImage(fp)

        dim = len(img.GetSize())
        if dim == 3:
            # 改变大小
            if new_size is not None:
                img = resize_image_itk(itkimage=img, newSize=new_size)
            img = sitk.GetArrayFromImage(img)  # 获取数据，获取到的数据格式是np.ndarray
            bs_list.append(img)
        elif dim == 4:
            new_img_list = list()
            for i in range(img.GetSize()[-1]):
                t_img = img[:, :, :, i]
                # 改变大小
                if new_size is not None:
                    t_img = resize_image_itk(itkimage=t_img, newSize=new_size)
                new_img = sitk.GetArrayFromImage(t_img)  # 获取数据，获取到的数据格式是np.ndarray
                new_img_list.append(new_img)
            img = np.array(new_img_list)
            img = np.transpose(img, (1, 2, 3, 0))

            if time_points_num is not None:
                img = img[:, :, :, :time_points_num]
            # img = img.astype(np.float32)
            # print(img.shape, img.dtype)
            bs_list.append(img)
            label_list.append(y)
    print('')
    return bs_list, label_list


# 找出被试中最短时间点的长度，并将所有被试时间点统一成这个长度
def same_time_point(roiss: list, min_len=None):
    if min_len is None:
        for i in range(len(roiss)):
            if min_len is None:
                min_len = roiss[i].shape[1]
                continue
            if min_len > roiss[i].shape[1]:
                min_len = roiss[i].shape[1]
    roiss = [rois[:, : min_len] for rois in roiss]  # 统一所有的被试的时间点长度
    return roiss, min_len


# 获取 皮尔逊系数矩阵组 (Tensor)
def getPearson_tensor(data, no_self_corr=False):
    pcorr = None
    is_batch = False
    # 判断是否为批次输入
    if type(data) == list:
        is_batch = True
    elif len(data.shape) > 2:
        is_batch = True
    if not is_batch:
        data = torch.unsqueeze(data, dim=0)
    for i in range(len(data)):
        d = data[i]
        if d.shape[0] == 1:
            p = torch.tensor([[1]], dtype=torch.float)
        else:
            p = torch.corrcoef(d)
        # 将NaN处理成0
        p = Nan2Zero(p)
        if no_self_corr is True:
            p = p - torch.eye(p.shape[0])
        p = torch.unsqueeze(p, dim=0)
        if pcorr is None:
            pcorr = p
        else:
            pcorr = torch.cat([pcorr, p], dim=0)
    if not is_batch:
        pcorr = torch.squeeze(pcorr, dim=0)
    return pcorr


def _pearson(data):
    assert len(data.shape) == 2  # 断言: 输入维度必须为(特征数, 特征长度)
    feature_num = data.shape[0]
    pearson_matrix = np.eye(N=feature_num, dtype=np.float32)
    p_matrix = np.zeros(shape=(feature_num, feature_num), dtype=np.float32)
    x, y = np.triu_indices(n=feature_num, k=1)
    for i in range(x.shape[0]):
        res = stats.pearsonr(data[x[i]], data[y[i]])
        pcorr = res[0]
        p = res[1]
        # 记录皮尔逊相关
        pearson_matrix[x[i], y[i]] = pcorr
        pearson_matrix[y[i], x[i]] = pcorr
        # 记录对应的p值
        p_matrix[x[i], y[i]] = p
        p_matrix[y[i], x[i]] = p

    return pearson_matrix, p_matrix


# 获取 皮尔逊系数矩阵组 (ndarray)
def getPearson_ndarray(data, no_self_corr=False, show=True, info=None):
    pcorr = None
    is_batch = False
    # 判断是否为批次输入
    if len(data.shape) > 2:
        is_batch = True
    if not is_batch:
        data = np.expand_dims(data, axis=0)
    for i in range(data.shape[0]):
        if show:
            print("\r- 正在计算第{}个数据的皮尔逊相关矩阵".format(i + 1), end='')
            if info is not None:
                print(f" - <{info}>", end='')
        d = data[i]
        if d.shape[0] == 1:
            p = np.array([[1.]], dtype=np.float32)
        else:
            p = np.corrcoef(d)
        # 将NaN处理成0
        p = Nan2Zero(p)
        if no_self_corr is True:
            p = p - np.eye(p.shape[0])
        p = np.expand_dims(p, axis=0)
        if pcorr is None:
            pcorr = p
        else:
            pcorr = np.concatenate([pcorr, p], axis=0)
    if show:
        print("")
    if not is_batch:
        pcorr = np.squeeze(pcorr, axis=0)

    return pcorr


# 获得有方向的皮尔逊矩阵
def get_direct_pearson_from_roi(rois1, rois2):
    p1, p2 = None, None
    # 批量输入判断
    if len(rois1.shape) > 2 or len(rois2.shape) > 2:
        assert len(rois1.shape) > 2 and len(rois2.shape) > 2
        assert rois1.shape[0] == rois2.shape[0]
        for i in range(rois1.shape[0]):
            n1 = rois1.shape[1]
            rois = np.concatenate([rois1[i], rois2[i]], axis=0) if type(rois1) is np.ndarray else torch.cat(
                [rois1[i], rois2[i]], dim=0)

            p = np.corrcoef(rois) if type(rois) is np.ndarray else torch.corrcoef(rois)
            p = Nan2Zero(p)

            t_p1 = p[0: n1, n1:]
            t_p2 = p[n1:, 0: n1]
            t_p1 = np.expand_dims(t_p1, axis=0) if type(t_p1) is np.ndarray else torch.unsqueeze(t_p1, dim=0)
            t_p2 = np.expand_dims(t_p2, axis=0) if type(t_p2) is np.ndarray else torch.unsqueeze(t_p2, dim=0)
            if p1 is None:
                p1 = t_p1
            else:
                p1 = np.concatenate([p1, t_p1], axis=0) if type(p1) is np.ndarray else torch.cat(
                    [p1, t_p1], dim=0)
            if p2 is None:
                p2 = t_p2
            else:
                p2 = np.concatenate([p2, t_p2], axis=0) if type(p2) is np.ndarray else torch.cat(
                    [p2, t_p2], dim=0)
    else:
        n1 = rois1.shape[0]
        rois = np.concatenate([rois1, rois2], axis=0) if type(rois1) is np.ndarray else torch.cat([rois1, rois2], dim=0)
        p = np.corrcoef(rois) if type(rois) is np.ndarray else torch.corrcoef(rois)

        p1 = p[0: n1, n1:]
        p2 = p[n1:, 0: n1]

    # 矫正皮尔逊中的无穷值
    p1 = np.where(p1 <= -np.inf, -1, p1)
    p1 = np.where(p1 >= np.inf, 1, p1)
    p2 = np.where(p2 <= -np.inf, -1, p2)
    p2 = np.where(p2 >= np.inf, 1, p2)

    return p1, p2


# 统计有多少NaN
def countNan(data):
    if not type(data) is np.ndarray:
        data = data.cpu().detach().numpy()
    unique, count = np.unique(np.isnan(data), return_counts=True)
    if True in unique:
        return count[unique.tolist().index(True)]
    else:
        return 0

# NaN->0
def Nan2Zero(data):
    zero = float(np.random.normal(loc=0, scale=1e-32, size=1)[0])
    if type(data) is np.ndarray:
        data[np.isnan(data)] = zero
    else:
        data[torch.isnan(data)] = zero

    return data


# 按标签种类分离样本
def apartByLabels(data, labels):
    # 通过标签获取样本种类
    classes = list(np.unique(labels))
    res = list()
    for c in classes:
        # 记录各种类在标签中的下标情况
        classes_state = (labels == c)
        # 通过各种类的下标情况取出对应的样本
        res.append(data[classes_state])

    return res


# 通过双样本t检验，返回合格的特征对应的下标
def TwoSampleTTest(features, labels):
    # 将正负样本分离
    res = apartByLabels(features, labels)
    assert len(res) == 2  # 断言：只能有2类样本
    s1, s2 = res[0], res[1]
    assert len(s1.shape) == 2 and len(s2.shape) == 2  # 断言：输入尺寸只能为2维 (样本数, 特征数)
    # 进行 双样本t检验
    t, p_values = stats.ttest_ind(s1, s2, axis=0, equal_var=False)
    # 获取排序后的下标(升序)
    f_idx = np.argsort(p_values)
    # 将p值排序
    p_values.sort()
    # 获取小于0.05的p值
    ok_state = p_values < 0.05
    # 获取经过 双样本t检验 合格的特征下标
    ok_f_idx = f_idx[ok_state]
    return ok_f_idx


# 通过1维特征的下标，定位原皮尔逊相关矩阵的位置(上三角中)
def get_idx_feature2matrix(feature_idx, naoqu_num: int):
    x, y = np.triu_indices(naoqu_num)
    if type(feature_idx) is int:  # 单个
        return (x[feature_idx], y[feature_idx])
    else:  # 批量
        res = list()
        for fi in feature_idx:
            res.append((x[fi], y[fi]))
        return res


# F-Score
def fscore_core(np,nn,xb,xbp,xbn,xkp,xkn):
    def sigmap (i,np,xbp,xkp):
        np=int(np)
        return sum([(xkp[i][k]-xbp[i])**2 for k in range(np)])

    def sigman (i,nn,xbn,xkn):
        nn=int(nn)
        return sum([(xkn[i][k]-xbn[i])**2 for k in range(nn)])
    n_feature = len(xb)
    fscores = []
    for i in range(n_feature):
        fscore_numerator = (xbp[i]-xb[i])**2 + (xbn[i]-xb[i])**2
        fscore_denominator = (1/float(np-1))*(sigmap(i,np,xbp,xkp))+ \
                             (1/float(nn-1))*(sigman(i,nn,xbn,xkn))
        fscores.append(fscore_numerator/fscore_denominator)

    return fscores


def _fscore(feature,classindex):
    n_instance = len(feature)
    n_feature = len(feature[0])
    np = sum(classindex)
    nn = n_instance - np
    xkp =[];xkn =[];xbp =[];xbn =[];xb=[]
    for i in range(n_feature):
        xkp_i = [];xkn_i = []
        for k in range(n_instance):
            if classindex[k] == 1:
                xkp_i.append(feature[k][i])
            else:
                xkn_i.append(feature[k][i])
        xkp.append(xkp_i)
        xkn.append(xkn_i)
        sum_xkp_i = sum(xkp_i)
        sum_xkn_i = sum(xkn_i)
        xbp.append(sum_xkp_i/float(np))
        xbn.append(sum_xkn_i/float(nn))
        xb.append((sum_xkp_i+sum_xkn_i)/float(n_instance))
    return fscore_core(np,nn,xb,xbp,xbn,xkp,xkn)


# 计算特征的F-Score
def FScore(features, labels, topk: int=-1):
    # F-Score分数
    fs = np.array(_fscore(features, labels))
    # 按降序排列的F-Score对应的下标
    ds_sort_idx = np.argsort(fs)[::-1]
    if topk != -1:
        fs = fs[:topk]
        ds_sort_idx = ds_sort_idx[:topk]
    return ds_sort_idx, fs


# 滑动窗口
def slice_window_tensor(data, window_width: int=40, stride: int=2, dim: int=-1):
    assert not (window_width <= 0 or stride <= 0)  # 断言: 不允许 window_width <= 0 或 stride <= 0
    assert not (window_width == 1 and stride == 1)  # 断言: 不允许 window_width 和 stride 同时为1
    # 把要切片的维度放到第0号下标上
    data = torch.transpose(data, 0, dim)
    cur_begin = 0
    cur_end = cur_begin + window_width
    slice_result = None
    while(True):
        # 判断是否可以切片
        if cur_end > data.shape[0]:
            break
        # 切片并还原维度位置, 然后在dim=1处扩展一个切片维度
        temp = torch.unsqueeze(torch.transpose(data[cur_begin: cur_end], 0, dim), dim=1)
        if slice_result is None:
            slice_result = temp
        else:
            slice_result = torch.cat([slice_result, temp], dim=1)
        # 更新前后位置
        cur_begin += stride
        cur_end = cur_begin + window_width

    # 尺寸: (被试数, 脑区数, 时间点数) -> (被试数, 切片数, 脑区数, 时间点数)
    # 尺寸: (被试数, 特征数) -> (被试数, 切片数, 特征数)
    return slice_result


# N4偏置场校正
def N4BiasField(imagePath: str):
    input_image = sitk.ReadImage(imagePath)
    mask_image = sitk.OtsuThreshold(input_image, 0, 1, 200)
    input_image = sitk.Cast(input_image, sitk.sitkFloat32)
    corrector = sitk.N4BiasFieldCorrectionImageFilter()
    output_image = corrector.Execute(input_image, mask_image)
    output_image = sitk.Cast(output_image, sitk.sitkInt16)
    # sitk.WriteImage(output_image, 'C:/Users/RONG/Desktop/Images/001/Lung_N4.nii.gz')
    # sitk.WriteImage(mask_image, 'C:/Users/RONG/Desktop/Images/001/Lung_mask.nii.gz')

    output_image = sitk.GetArrayFromImage(output_image)
    mask_image = sitk.GetArrayFromImage(mask_image)

    return output_image, mask_image


# 生成.doc .docx文件
# https://blog.csdn.net/qq_37746855/article/details/115271750
def makeDoc(text: str, save_path=None):
    from docx import Document
    from docx.shared import Pt

    if save_path is None:
        save_path = './new.docx'

    # 简单的打开word，输入数据，关闭word
    document = Document()
    # 设置word字体大小
    style = document.styles['Normal']
    font = style.font
    font.size = Pt(8)
    # 向word里增加段落
    document.add_paragraph(text)
    document.save(save_path)

    # # 在一个段落中增加文字
    # document = Document()
    # paragraph = document.add_paragraph('Hello, ')
    # # 增加文字
    # paragraph.add_run('tgenkidu')


# 执行双样本t检验
def ttest_t_p(A, B):
    t, p = stats.ttest_ind(A, B)
    return t, p


# 复制文件
def copy_file(src_path, new_path):
    shutil.copyfile(src_path, new_path)


# 数据归一化
def normalization(data):
    try:
        _range = np.max(data) - np.min(data)
        return (data - np.min(data)) / _range
    except TypeError:
        _range = torch.max(data) - torch.min(data)
        return (data - torch.min(data)) / _range


# 数据标准化
def standardization(data, axis_dims=0):
    try:
        mu = np.mean(data, axis=axis_dims)
        sigma = np.std(data, axis=axis_dims)
    except TypeError:
        mu = torch.mean(data, dim=axis_dims)
        sigma = torch.std(data, dim=axis_dims)
    return (data - mu) / sigma


# 多线程 计算 皮尔逊系数矩阵组 (ndarray)
thread_info = None
pcorr_dict = dict()
pcorr_idx_list = None
# 计算 皮尔逊系数矩阵组 (ndarray)
def getPearson_list(data, no_self_corr=False, show=True, info=None, workers_num=9):
    global thread_info
    global pcorr_dict
    global pcorr_idx_list

    thread_info = [0 for _ in range(workers_num)]
    pcorr_idx_list = [i for i in range(len(data))]


    import concurrent.futures
    def _pcorr(data, idx_list, thread_id):
        global thread_info

        for idx in idx_list:
            thread_info[thread_id] += 1
            d = data[idx]
            if d.shape[0] == 1:
                p = np.array([[1.]], dtype=np.float32) if type(data[0]) == np.ndarray else torch.from_numpy(np.array([[1.]], dtype=np.float32))
            else:
                p = np.corrcoef(d) if type(data[0]) == np.ndarray else torch.corrcoef(d)
            # 将NaN处理成0
            p = Nan2Zero(p)
            if show:
                print(f"\r- 正在计算皮尔逊相关矩阵: {thread_info} -> ({np.array(thread_info).sum()}/{len(data)})", end="")
            pcorr_dict[idx] = p
            pass

    # 多线程 计算皮尔逊相关矩阵
    per_len = math.ceil(len(data) / workers_num)
    # 分割数据
    cur_idx = 0
    tq = [list() for _ in range(workers_num)]
    for idx in pcorr_idx_list:
        tq[cur_idx // per_len].append(idx)
        cur_idx += 1
    # 开启多线程
    with concurrent.futures.ThreadPoolExecutor() as executor:
        for i in range(len(tq)):
            executor.submit(lambda a: _pcorr(*a), (data, tq[i], i))
        executor.shutdown()
        if show:
            print("")
    # 合并结果
    res_pcorr = list()
    for i in range(len(pcorr_dict.keys())):
        if type(data[0]) == np.ndarray:
            if type(pcorr_dict[i]) != np.ndarray:
                res_pcorr.append(pcorr_dict[i].numpy())
            else:
                res_pcorr.append(pcorr_dict[i])
        else:
            if type(pcorr_dict[i]) != torch.Tensor:
                res_pcorr.append(pcorr_dict[i])
            else:
                res_pcorr.append(pcorr_dict[i].numpy())

    pcorr = np.array(res_pcorr) if type(data[0]) == np.ndarray else torch.from_numpy(np.array(res_pcorr))

    # 初始化全局参数
    thread_info = None
    pcorr_dict = dict()
    pcorr_idx_list = None
    
    return pcorr